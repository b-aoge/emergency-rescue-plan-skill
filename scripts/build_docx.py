# -*- coding: utf-8 -*-
"""从 plan_part1~5.md 重建《施工现场应急救援预案》docx（无重复），套用排版 + 目录真实页码。

用法：
  python build_docx.py
前置：当前工作目录含 plan_part1.md ~ plan_part5.md；若存在 toc_pages.json 则目录写死静态页码。
输出：施工现场应急救援预案_生成版.docx
"""
import re, os
from docx import Document
from docx.shared import Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_UNDERLINE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK = "宋体"
HEI = "黑体"

SRC = ["plan_part1.md", "plan_part2.md", "plan_part3.md", "plan_part4.md", "plan_part5.md"]
OUT = "施工现场应急救援预案_生成版.docx"

# 目录条目（与文档中实际标题一致；用于书签匹配与页码域）
TOC = [
    "一、编制依据", "二、基本原则与方针", "三、工程项目基本情况",
    "四、可能发生事故的确定及主要危险源", "五、应急机构的组成、责任和分工",
    "六、报警信号与通讯", "七、应急响应与救援", "八、有关规定和要求", "九、附件",
    "附件1 常见事故急救常识", "附件2 火灾逃生自救知识", "附件3 灭火器的使用常识",
    "附件4 各类事故的预防措施", "附件5 施工总平面布置图",
    "附件6 应急救援路线图",
]
TOC_IDX = {re.sub(r"\s+", " ", t).strip(): i for i, t in enumerate(TOC)}

# 若已有 Word 计算好的真实页码，则直接写死为静态文本（避免在线预览不刷新 PAGEREF 域导致页码全为占位值）
# 优先读取当前工作目录的 toc_pages.json，找不到再回退到脚本所在目录。
try:
    import json as _json
    _cand = [os.path.join(os.getcwd(), "toc_pages.json"),
             os.path.join(os.path.dirname(os.path.abspath(__file__)), "toc_pages.json")]
    TOC_PAGES = {}
    for _p in _cand:
        if os.path.exists(_p):
            with open(_p, encoding="utf-8") as _f:
                TOC_PAGES = _json.load(_f)
            break
except Exception:
    TOC_PAGES = {}

def norm(s):
    return re.sub(r"\s+", " ", s).strip()

# ---------- 字体/排版辅助 ----------
def set_cjk(run, cjk, size_pt, bold=False, color=None):
    run.font.name = cjk
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = color if color is not None else RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), cjk)
    rf.set(qn("w:ascii"), cjk)
    rf.set(qn("w:hAnsi"), cjk)

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "000000")
        borders.append(e)
    tblPr.append(borders)

def add_bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "999999")
    pbdr.append(bottom); pPr.append(pbdr)

def render_cover(cover_lines):
    """封面版式：标题居中 + 签名区页面中下部 + 底部编制单位/日期居中。"""
    title_lines, sign_lines, bottom_lines = [], [], []
    stage = "top"
    for ln in cover_lines:
        s = ln.rstrip()
        if not s.strip():
            continue
        if s.startswith(("编制：", "审核：", "批准：")):
            stage = "sign"
        elif s.startswith(("编制单位：", "编制日期：")):
            stage = "bottom"
        elif stage == "sign" and not s.startswith(("编制：", "审核：", "批准：")):
            stage = "bottom"
        if stage == "top":
            title_lines.append(s)
        elif stage == "sign":
            sign_lines.append(s)
        else:
            bottom_lines.append(s)

    for s in title_lines:
        if s.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(s[2:].strip()); set_cjk(r, HEI, 22, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(12)
        elif s.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(s[3:].strip()); set_cjk(r, HEI, 16, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(180)
        else:
            add_body(s)

    for s in sign_lines:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run(s[:3]); set_cjk(r1, CJK, 12)
        r2 = p.add_run(" " * 22); set_cjk(r2, CJK, 12); r2.font.underline = WD_UNDERLINE.SINGLE

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(28)

    for s in bottom_lines:
        if s.startswith(("编制单位：", "编制日期：")):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
            r = p.add_run(s); set_cjk(r, CJK, 12)
        else:
            add_body(s)

def add_field(paragraph, instr):
    r = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = " " + instr + " "
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r._r.append(b); r._r.append(t); r._r.append(e)

_bid = [1000]
def add_bookmark(paragraph, name):
    _bid[0] += 1
    p = paragraph._p
    bs = OxmlElement("w:bookmarkStart"); bs.set(qn("w:id"), str(_bid[0])); bs.set(qn("w:name"), name)
    be = OxmlElement("w:bookmarkEnd"); be.set(qn("w:id"), str(_bid[0]))
    p.insert(0, bs); p.append(be)

def add_xml_flag(p, tag):
    pPr = p._p.get_or_add_pPr()
    el = pPr.find(qn(tag))
    if el is None:
        pPr.append(OxmlElement(tag))

# ---------- 块渲染 ----------
doc = Document()
for sn in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
    try:
        st = doc.styles[sn]
    except KeyError:
        continue
    st.font.name = CJK; st.font.size = Pt(10.5); st.font.color.rgb = RGBColor(0, 0, 0)
    st.element.rPr.get_or_add_rFonts().set(qn("w:eastAsia"), CJK)
    st.element.rPr.get_or_add_rFonts().set(qn("w:ascii"), CJK)
    st.element.rPr.get_or_add_rFonts().set(qn("w:hAnsi"), CJK)

in_cover = True
first_h1 = True
toc_done = False
LIST_RE = re.compile(r"^(\d+[．.、]|[（(]\d+[)）]|[（(][一二三四五六七八九十]+[)）]|[一二三四五六七八九十]+[．.、])")

def add_heading(text, level):
    global first_h1
    style_name = "Heading %d" % level if level <= 3 else "Normal"
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if level == 1:
        if first_h1:
            set_cjk(run, HEI, 22, bold=True); first_h1 = False
            p.paragraph_format.space_before = Pt(0)
        else:
            set_cjk(run, HEI, 16, bold=True)
            p.paragraph_format.space_before = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
    elif level == 2:
        set_cjk(run, HEI, 14, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if in_cover else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
    else:
        set_cjk(run, HEI, 12, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    # 附件6 路线图单独成页，保证标题与图同页
    if text.startswith("附件6"):
        add_xml_flag(p, "w:keepNext"); add_xml_flag(p, "w:pageBreakBefore")
    key = norm(text)
    if key in TOC_IDX:
        add_bookmark(p, "_Toc%d" % TOC_IDX[key])

def add_divider():
    global in_cover
    p = doc.add_paragraph(); add_bottom_border(p)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(6)
    in_cover = False

def insert_toc():
    ph = doc.add_paragraph()
    r = ph.add_run("目　录"); set_cjk(r, HEI, 16, bold=True)
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph.paragraph_format.space_after = Pt(8); ph.paragraph_format.line_spacing = 1.25
    sec = doc.sections[0]
    avail = sec.page_width - sec.left_margin - sec.right_margin
    for i, title in enumerate(TOC):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25; p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(avail, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r0 = p.add_run(title); set_cjk(r0, CJK, 10.5)
        p.add_run("\t")
        num = TOC_PAGES.get(str(i))
        if num:
            r1 = p.add_run(str(num)); set_cjk(r1, CJK, 10.5)   # 静态页码，所有阅读器均可读
        else:
            add_field(p, "PAGEREF _Toc%d" % i); set_cjk(p.runs[-1], CJK, 10.5)
    doc.add_page_break()

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text); set_cjk(run, CJK, 10.5)
    p.paragraph_format.line_spacing = 1.25; p.paragraph_format.space_after = Pt(3)
    if in_cover:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif LIST_RE.match(text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(21)

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text); set_cjk(run, CJK, 10.5)
    p.paragraph_format.line_spacing = 1.25; p.paragraph_format.space_after = Pt(2)

def set_cell_vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign"); v.set(qn("w:val"), "center"); tcPr.append(v)

def set_cell_margins(table, top=40, bottom=40, left=90, right=90):
    tblPr = table._tbl.tblPr
    ex = tblPr.find(qn("w:tblCellMar"))
    if ex is not None: tblPr.remove(ex)
    m = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{tag}"); e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa"); m.append(e)
    tblPr.append(m)

def add_image(alt, path):
    """插入图片：平面图最多 11cm，手机长截图（路线图）最多 8cm，居中并与标题同页。"""
    full = path if os.path.isabs(path) else os.path.join(os.getcwd(), path)
    if not os.path.exists(full):
        p = doc.add_paragraph(); r = p.add_run(f"[图片未找到：{full}]"); set_cjk(r, CJK, 10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; return
    sec = doc.sections[0]
    avail = sec.page_width - sec.left_margin - sec.right_margin
    default_w = Cm(8) if "路线" in alt else Cm(11)
    max_w = min(avail, default_w)
    doc.add_picture(full, width=max_w)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_paragraph.paragraph_format.keep_with_next = True
    last_paragraph.paragraph_format.space_after = Pt(6)

def add_table(tbl_lines):
    rows = [ [c.strip() for c in tl.strip().strip("|").split("|")] for tl in tbl_lines ]
    header = rows[0]; ncol = len(header)
    sec = doc.sections[0]
    avail = int(sec.page_width - sec.left_margin - sec.right_margin)
    t = doc.add_table(rows=len(rows), cols=ncol); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tblPr = t._tbl.tblPr
    old_tw = tblPr.find(qn("w:tblW"))
    if old_tw is not None: tblPr.remove(old_tw)
    tw = OxmlElement("w:tblW"); tw.set(qn("w:w"), "5000"); tw.set(qn("w:type"), "pct"); tblPr.append(tw)
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); tblPr.append(lay)
    set_table_borders(t); set_cell_margins(t)
    RATIO = {
        ("姓名", "职务", "联系电话"): [1.4, 3.0, 2.6],
        ("单位或姓名", "电话"): [3.0, 1.2],
        ("序号", "事故类别", "事故原因", "现场救援措施", "演练时间"): [0.6, 1.7, 2.6, 3.6, 1.0],
    }
    ratios = RATIO.get(tuple(header), [1.0] * ncol)
    tot = sum(ratios)
    colw = [int(avail * r / tot) for r in ratios]
    for ci in range(ncol):
        t.columns[ci].width = Emu(colw[ci])
    for r, row in enumerate(rows):
        for c in range(ncol):
            cell = t.cell(r, c); cell.width = Emu(colw[c]); set_cell_vcenter(cell)
            cp = cell.paragraphs[0]; cp.text = ""
            run = cp.add_run(row[c] if c < len(row) else ""); set_cjk(run, CJK, 10)
            cp.paragraph_format.line_spacing = 1.15; cp.paragraph_format.space_after = Pt(0)
            if r == 0:
                run.bold = True; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER; shade_cell(cell, "D9E2F3")
            else:
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if (c == 0 or c == ncol - 1) else WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

sp0 = doc.add_paragraph(); sp0.paragraph_format.space_after = Pt(18)

# ---------- 解析并渲染 ----------
text = ""
for f in SRC:
    text += open(f, encoding="utf-8").read() + "\n"
lines = text.split("\n"); N = len(lines); i = 0

if lines and lines[0].strip() == "<!-- cover -->":
    end = 0
    for idx, ln in enumerate(lines):
        if ln.strip() == "<!-- /cover -->":
            end = idx; break
    if end:
        render_cover(lines[1:end]); doc.add_page_break(); first_h1 = False
        lines = lines[end + 1:]; N = len(lines); i = 0

while i < N:
    s = lines[i].strip()
    if not s:
        i += 1; continue
    if s == "---":
        if not toc_done:
            insert_toc(); toc_done = True
        else:
            add_divider()
        i += 1; continue
    if s == "<!-- pagebreak -->":
        doc.add_page_break(); i += 1; continue
    if s.startswith("#"):
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        add_heading(m.group(2).strip(), len(m.group(1))); i += 1; continue
    if s.startswith("|"):
        tbl = []
        while i < N and lines[i].strip().startswith("|"):
            tbl.append(lines[i].strip()); i += 1
        add_table(tbl); continue
    img = re.match(r"^!\[(.*?)\]\((.*?)\)$", s)
    if img:
        add_image(img.group(1), img.group(2)); i += 1; continue
    if s.startswith("- "):
        add_bullet(s[2:].strip()); i += 1; continue
    add_body(s); i += 1

# ---------- 页码（页脚居中） ----------
section = doc.sections[0]
fp = section.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = fp.add_run("第 "); set_cjk(r1, CJK, 10.5)
fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
r2 = fp.add_run(); r2._r.append(fld1); r2._r.append(instr); r2._r.append(fld2)
r3 = fp.add_run(" 页"); set_cjk(r3, CJK, 10.5)

doc.save(OUT)

# 设置 Word 打开时自动更新域（目录/页码）
def enable_update_fields(docx_path):
    import zipfile, tempfile, shutil
    settings_path = "word/settings.xml"
    tmp_dir = tempfile.mkdtemp()
    tmp_docx = os.path.join(tmp_dir, os.path.basename(docx_path))
    shutil.copy2(docx_path, tmp_docx)
    with zipfile.ZipFile(tmp_docx, 'r') as zin:
        try:
            settings_xml = zin.read(settings_path).decode('utf-8')
        except KeyError:
            settings_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:settings>'
    if 'w:updateFields' not in settings_xml:
        settings_xml = settings_xml.replace('</w:settings>', '<w:updateFields w:val="true"/></w:settings>')
    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        with zipfile.ZipFile(tmp_docx, 'r') as zin:
            for item in zin.namelist():
                if item == settings_path:
                    zout.writestr(item, settings_xml)
                else:
                    zout.writestr(item, zin.read(item))
    shutil.rmtree(tmp_dir)

enable_update_fields(OUT)
print("SAVED:", os.path.abspath(OUT), "size:", os.path.getsize(OUT))
